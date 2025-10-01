import requests
import os
import uuid
from docx import Document
from dotenv import load_dotenv
from bs4 import BeautifulSoup
from langchain_openai.chat_models.azure import AzureChatOpenAI
load_dotenv()

azure_endpoint = os.getenv('azure_endpoint')
api_version_azure = os.getenv('api_version_azure')
api_key_azure = os.getenv('api_key_azure')
deployment_name = os.getenv( "deployment_name")

subscription_key = os.getenv('subscription_key')
endpoint = os.getenv('endpoint')
location = os.getenv('location')

target_language = 'pt-br'

def translator_text(text, target_language):
  path = '/translate'
  constructed_url = endpoint + path
  headers = {
    'Ocp-Apim-Subscription-Key': subscription_key,
    'Ocp-Apim-Subscription-Region': location,
    'Content-type': 'application/json',
    'X-ClientTraceId': str(os.urandom(16))
  }

  body =[{
    'text': text
  }]
  
  params = {
    'api-version': '3.0',
    'from': 'en',
    'to': target_language
  }

  request = requests.post(constructed_url, params=params, headers=headers, json=body)
  
  response = request.json()
  
  return response[0]["translations"][0]["text"]


def translate_document(path):
  document = Document(path)
  full_text = []
  for paragraph in document.paragraphs:
    if paragraph.text.strip():
      translate_text = translator_text(paragraph.text, target_language)
      full_text.append(translate_text)
    else:
      full_text.append("")
  
  translate_doc = Document()
  for line in full_text:
    translate_doc.add_paragraph(line) 
   
  path_translated = path.replace(".docx", f"_{target_language}.docx")
  translate_doc.save(path_translated)
  return path_translated


def extract_text_from_url(url):
  response = requests.get(url)
  
  if response.status_code == 200:
    soup = BeautifulSoup(response.text, 'html.parser')
    for scripy_or_style in soup(['script', 'style']):
      scripy_or_style.decompose()
    texto = soup.get_text(separator= ' ')
    
    linhas = (line.strip() for line in texto.splitlines())
    parts = (phrase.strip() for line in linhas for phrase in line.split(" "))
    texto_limpo = '\n'.join(part for part in parts if part)
    
    return texto_limpo
  else:
      print(f"Failed to fetch the URL. Status code: {response.status_code}")
      return None


client = AzureChatOpenAI(
  azure_endpoint = azure_endpoint,
  api_version = api_version_azure,
  api_key = api_key_azure,
  deployment_name = deployment_name,
  max_retries=0
)

def translate_article(text, target_language):
  messages = [
    ('system', "Você atua como tradutor de texto"),
    ('user', f"traduza o {text} para o idioma {target_language}, e responda em markdown")
  ]
  response = client.invoke(messages)
  print(response.content)
  return(response.content)

url = "https://dev.to/kenakamu/azure-open-ai-in-vnet-3alo"

text = extract_text_from_url(url)
article = translate_article(text, target_language)

print(article)


